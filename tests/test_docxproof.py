from __future__ import annotations

import json
import re
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

from docxproof import engine
from docxproof.providers import _resolve_deepseek_reasoning_effort
from docxproof.schemas import Correction, CorrectionBatch, ProposedPatch
from docxproof.settings import (
    EDITABLE_END,
    EDITABLE_START,
    resolve_model,
    resolve_provider_and_model,
)


class FakeAdapter:
    provider = "fake"
    model = "fake-proofreader"
    retryable_exceptions = ()

    def __init__(self) -> None:
        self.calls: list[str] = []

    def complete(self, system_prompt: str, user_prompt: str) -> CorrectionBatch:
        self.calls.append(user_prompt)
        corrections = []
        for old, new, reason in [
            ("Their", "They're", "their/they are grammatical error"),
            ("are wrong", "is wrong", "subject-verb agreement"),
            ("A apple", "An apple", "article agreement"),
            ("It work", "It works", "subject-verb agreement"),
        ]:
            pos = user_prompt.find(old)
            if pos < 0:
                continue
            editable_start = user_prompt.find(EDITABLE_START)
            editable_end = user_prompt.find(EDITABLE_END)
            if not (editable_start <= pos < editable_end):
                continue
            labels = list(re.finditer(r"\[\[([^\]]+)\]\]", user_prompt[:pos]))
            if labels:
                corrections.append(
                    Correction(
                        paragraph_id=labels[-1].group(1),
                        original=old,
                        replacement=new,
                        occurrence=1,
                        reason=reason,
                    )
                )
        return CorrectionBatch(corrections=corrections)


class DocxProofTests(unittest.TestCase):
    def _word_story(self, count: int = 1000):
        temp_dir = tempfile.TemporaryDirectory()
        path = Path(temp_dir.name) / "window.docx"
        doc = Document()
        doc.add_paragraph(" ".join(f"w{i}" for i in range(count)))
        doc.save(path)
        _, stories = engine.load_docx_stories(path, ("main",))
        return temp_dir, stories[0]

    def test_model_choices_are_deliberately_limited(self) -> None:
        self.assertEqual(resolve_model("openai", None), "gpt-5.6-terra")
        self.assertEqual(resolve_model("deepseek", None), "deepseek-v4-pro")
        with self.assertRaises(ValueError):
            resolve_model("openai", "arbitrary-model")

    def test_provider_can_be_inferred_from_model(self) -> None:
        self.assertEqual(
            resolve_provider_and_model(None, "deepseek-v4-pro"),
            ("deepseek", "deepseek-v4-pro"),
        )
        self.assertEqual(
            resolve_provider_and_model(None, "gpt-5.6"),
            ("openai", "gpt-5.6"),
        )
        with self.assertRaises(ValueError):
            resolve_provider_and_model("openai", "deepseek-v4-pro")

    def test_deepseek_flash_uses_low_reasoning_by_default(self) -> None:
        self.assertEqual(
            _resolve_deepseek_reasoning_effort("deepseek-v4-flash", "medium"),
            "low",
        )
        self.assertEqual(
            _resolve_deepseek_reasoning_effort("deepseek-v4-pro", "medium"),
            "high",
        )
        self.assertEqual(
            _resolve_deepseek_reasoning_effort("deepseek-v4-pro", "none"),
            "none",
        )

    def test_windows_overlap_but_keep_primary_regions_for_reporting(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        windows = engine.make_windows(story, 400, 100)
        self.assertEqual(
            [(w.context_word_start, w.context_word_end) for w in windows],
            [(0, 400), (300, 700), (600, 1000)],
        )
        self.assertEqual(
            [(w.target_word_start, w.target_word_end) for w in windows],
            [(0, 350), (350, 650), (650, 1000)],
        )
        self.assertIn("w320", windows[0].rendered_text)
        self.assertIn("w320", windows[1].rendered_text)

    def test_identical_overlap_edits_are_merged_and_corroborated(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        first, second = engine.make_windows(story, 400, 100)[:2]
        paragraph_id = story.paragraphs[0].paragraph_id
        batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w320",
                    replacement="fixed",
                    reason="test correction",
                )
            ]
        )
        patches = []
        for window in (first, second):
            accepted, rejected = engine.validate_batch(story, window, batch)
            self.assertFalse(rejected)
            patches.extend(accepted)

        accepted, rejected = engine.resolve_conflicts(patches)
        self.assertFalse(rejected)
        self.assertEqual(len(accepted), 1)
        self.assertEqual(len(accepted[0].source_windows), 2)
        self.assertTrue(accepted[0].corroborated)
        self.assertTrue(accepted[0].primary_owner_seen)

    def test_equivalent_broad_and_minimal_edits_are_normalized_and_merged(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        first, second = engine.make_windows(story, 400, 100)[:2]
        paragraph_id = story.paragraphs[0].paragraph_id

        broad = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w320 w321",
                    replacement="fixed w321",
                    reason="broad form",
                )
            ]
        )
        minimal = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w320",
                    replacement="fixed",
                    reason="minimal form",
                )
            ]
        )
        patches = []
        patches.extend(engine.validate_batch(story, first, broad)[0])
        patches.extend(engine.validate_batch(story, second, minimal)[0])

        accepted, rejected = engine.resolve_conflicts(patches)
        self.assertFalse(rejected)
        self.assertEqual(len(accepted), 1)
        self.assertEqual(accepted[0].original, "w320")
        self.assertEqual(accepted[0].replacement, "fixed")
        self.assertEqual(len(accepted[0].source_windows), 2)

    def test_correction_crossing_primary_boundary_can_be_used(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        first, second = engine.make_windows(story, 400, 100)[:2]
        paragraph_id = story.paragraphs[0].paragraph_id
        batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w349 w350",
                    replacement="fixed phrase",
                    reason="cross-boundary correction",
                )
            ]
        )
        patches = []
        for window in (first, second):
            patches.extend(engine.validate_batch(story, window, batch)[0])

        accepted, rejected = engine.resolve_conflicts(patches)
        self.assertFalse(rejected)
        self.assertEqual(len(accepted), 1)
        self.assertFalse(accepted[0].primary_owner_seen)
        self.assertEqual(len(accepted[0].source_windows), 2)

    def test_different_overlap_replacements_use_preceding_context_precedence(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        first, second = engine.make_windows(story, 400, 100)[:2]
        paragraph_id = story.paragraphs[0].paragraph_id

        # w370 is outside the first window's nominal primary region but the
        # first window still has much more preceding context than the second.
        first_batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w370",
                    replacement="first-choice",
                    reason="earlier-window alternative",
                )
            ]
        )
        second_batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w370",
                    replacement="second-choice",
                    reason="later-window alternative",
                )
            ]
        )
        first_patch = engine.validate_batch(story, first, first_batch)[0][0]
        second_patch = engine.validate_batch(story, second, second_batch)[0][0]
        self.assertGreater(
            first_patch.preceding_context_words,
            second_patch.preceding_context_words,
        )
        self.assertFalse(first_patch.primary_owner_seen)
        self.assertTrue(second_patch.primary_owner_seen)

        accepted, rejected = engine.resolve_conflicts([first_patch, second_patch])
        self.assertEqual(len(accepted), 1)
        self.assertEqual(accepted[0].replacement, "first-choice")
        self.assertEqual(len(rejected), 1)
        self.assertIn("more preceding context", rejected[0].rejection)

    def test_later_window_can_add_non_overlapping_information(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        first, second = engine.make_windows(story, 400, 100)[:2]
        paragraph_id = story.paragraphs[0].paragraph_id

        first_batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w320",
                    replacement="alpha",
                    reason="first correction",
                )
            ]
        )
        second_batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w380",
                    replacement="beta",
                    reason="additional correction",
                )
            ]
        )
        patches = engine.validate_batch(story, first, first_batch)[0]
        patches += engine.validate_batch(story, second, second_batch)[0]

        accepted, rejected = engine.resolve_conflicts(patches)
        self.assertFalse(rejected)
        self.assertEqual({p.replacement for p in accepted}, {"alpha", "beta"})

    def test_lower_context_window_is_used_when_it_is_the_only_proposal(self) -> None:
        temp_dir, story = self._word_story()
        self.addCleanup(temp_dir.cleanup)
        second = engine.make_windows(story, 400, 100)[1]
        paragraph_id = story.paragraphs[0].paragraph_id
        batch = CorrectionBatch(
            corrections=[
                Correction(
                    paragraph_id=paragraph_id,
                    original="w320",
                    replacement="fixed",
                    reason="missed by the earlier window",
                )
            ]
        )
        patch = engine.validate_batch(story, second, batch)[0][0]
        accepted, rejected = engine.resolve_conflicts([patch])
        self.assertFalse(rejected)
        self.assertEqual(accepted, [patch])

    def test_lower_precedence_broad_conflict_does_not_cancel_separate_winners(self) -> None:
        def patch(
            start: int,
            end: int,
            replacement: str,
            *,
            preceding: int,
            window_index: int,
        ) -> ProposedPatch:
            return ProposedPatch(
                part_name="word/document.xml",
                window_key=f"w{window_index}",
                paragraph_id="D-P000001",
                original="x" * (end - start),
                replacement=replacement,
                occurrence=1,
                reason="test",
                local_start=start,
                local_end=end,
                global_start=start,
                global_end=end,
                source_windows=(f"w{window_index}",),
                primary_owner_seen=True,
                window_index=window_index,
                preceding_context_words=preceding,
                following_context_words=20,
            )

        left = patch(0, 5, "a", preceding=300, window_index=1)
        broad = patch(4, 9, "b", preceding=40, window_index=2)
        right = patch(8, 12, "c", preceding=300, window_index=1)
        accepted, rejected = engine.resolve_conflicts([left, broad, right])
        self.assertEqual({p.replacement for p in accepted}, {"a", "c"})
        self.assertEqual(len(rejected), 1)
        self.assertEqual(rejected[0].replacement, "b")

    def test_equal_precedence_conflict_is_left_unchanged(self) -> None:
        def patch(replacement: str) -> ProposedPatch:
            return ProposedPatch(
                part_name="word/document.xml",
                window_key="same-window",
                paragraph_id="D-P000001",
                original="wrong",
                replacement=replacement,
                occurrence=1,
                reason="test",
                local_start=10,
                local_end=15,
                global_start=10,
                global_end=15,
                source_windows=("same-window",),
                primary_owner_seen=True,
                window_index=1,
                preceding_context_words=100,
                following_context_words=100,
            )

        accepted, rejected = engine.resolve_conflicts([patch("one"), patch("two")])
        self.assertFalse(accepted)
        self.assertEqual(len(rejected), 2)
        self.assertTrue(all("equal-precedence" in p.rejection for p in rejected))

    def test_end_to_end_is_sequential_and_preserves_formatting_and_media(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "fixture.docx"
            output = root / "fixture_out.docx"
            report_txt = root / "fixture_out.proofreading.txt"
            report_json = root / "fixture_out.proofreading.json"
            image_path = root / "image.png"

            Image.new("RGB", (20, 20), "white").save(image_path)
            doc = Document()
            paragraph = doc.add_paragraph()
            bold = paragraph.add_run("Their")
            bold.bold = True
            italic = paragraph.add_run(" happy with the result.")
            italic.italic = True
            paragraph.add_run(" This sentence is already correct.")
            paragraph.add_run().add_picture(str(image_path))
            doc.add_table(rows=1, cols=1).cell(0, 0).text = "This are wrong."
            section = doc.sections[0]
            section.header.paragraphs[0].text = "A apple appears here."
            section.footer.paragraphs[0].text = "It work in the footer."
            doc.save(source)

            with zipfile.ZipFile(source) as archive:
                source_parts = {
                    item.filename: archive.read(item.filename)
                    for item in archive.infolist()
                }

            adapter = FakeAdapter()
            result = engine.run_proofreader(
                source,
                output,
                adapter,
                include=("main", "headers", "footers"),
                window_words=400,
                overlap_words=100,
                verify=False,
                report_txt_path=report_txt,
                report_json_path=report_json,
            )
            self.assertEqual(len(result.accepted), 4)
            self.assertEqual(len(adapter.calls), result.completed_windows)

            out_doc = Document(output)
            self.assertTrue(out_doc.paragraphs[0].text.startswith("They're happy"))
            self.assertTrue(out_doc.paragraphs[0].runs[0].bold)
            self.assertTrue(out_doc.paragraphs[0].runs[1].italic)
            self.assertEqual(out_doc.tables[0].cell(0, 0).text, "This is wrong.")
            self.assertEqual(
                out_doc.sections[0].header.paragraphs[0].text,
                "An apple appears here.",
            )
            self.assertEqual(
                out_doc.sections[0].footer.paragraphs[0].text,
                "It works in the footer.",
            )

            with zipfile.ZipFile(output) as archive:
                output_parts = {
                    item.filename: archive.read(item.filename)
                    for item in archive.infolist()
                }
            self.assertEqual(
                source_parts["word/media/image1.png"],
                output_parts["word/media/image1.png"],
            )
            self.assertEqual(source_parts["word/styles.xml"], output_parts["word/styles.xml"])
            payload = json.loads(report_json.read_text(encoding="utf-8"))
            self.assertTrue(payload["accepted"])
            self.assertIn("Applied corrections: 4", report_txt.read_text(encoding="utf-8"))

    def test_occurrence_is_scoped_to_visible_window(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "occurrence.docx"
            words = [f"w{i}" for i in range(800)]
            words[100] = "teh"
            words[500] = "teh"
            doc = Document()
            doc.add_paragraph(" ".join(words))
            doc.save(path)

            _, stories = engine.load_docx_stories(path, ("main",))
            story = stories[0]
            second = engine.make_windows(story, 400, 100)[1]
            paragraph_id = story.paragraphs[0].paragraph_id
            batch = CorrectionBatch(
                corrections=[
                    Correction(
                        paragraph_id=paragraph_id,
                        original="teh",
                        replacement="the",
                        occurrence=1,
                        reason="spelling",
                    )
                ]
            )
            accepted, rejected = engine.validate_batch(story, second, batch)
            self.assertFalse(rejected)
            self.assertEqual(len(accepted), 1)
            self.assertEqual(accepted[0].local_start, story.paragraphs[0].text.rfind("teh") + 1)


if __name__ == "__main__":
    unittest.main()
